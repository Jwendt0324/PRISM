#!/usr/bin/env python3
"""
Mainframe Full DOCX Converter
Converts ALL markdown files in the Mainframe to readable .docx files.
Outputs to ~/Desktop/Mainframe Docs/ organized by section.
Skips session logs (internal) unless --include-logs is passed.
"""

import os
import sys
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

MAINFRAME_DIR = os.environ.get("MAINFRAME_DIR", os.path.expanduser("~/Documents/Claude/Mainframe"))
OUTPUT_DIR = os.environ.get("OUTPUT_DIR", os.path.expanduser("~/Desktop/Mainframe Docs"))

# Map Mainframe subfolders to human-readable output folder names
FOLDER_MAP = {
    "sops/client-work": "SOPs — Client Work",
    "sops/business-ops": "SOPs — Business Ops",
    "sops/file-management": "SOPs — File Management",
    "sops/templates": "SOPs — Templates",
    "memory-bank": "Memory Bank",
    "team-ops": "Team Operating System",
    "content-audit": "Content Audit",
    "content-audit/drafts/articles": "Content Drafts — Articles",
    "content-audit/drafts/video-scripts": "Content Drafts — Video Scripts",
    "content-audit/drafts/linkedin-posts": "Content Drafts — LinkedIn Posts",
    "content-audit/drafts/email-snippets": "Content Drafts — Email Snippets",
    "claude-code": "Claude Code Prompts",
    "skills": "Skills",
    "logs": "Session Logs",
}

# Files/folders to skip
SKIP_PATTERNS = [
    "logs/",          # Session logs are internal
    "_readable/",     # Old readable copies
    "node_modules/",
    ".DS_Store",
    "convert_sops_to_docx.py",
    "scripts/",
]

INCLUDE_LOGS = "--include-logs" in sys.argv


def get_output_folder(rel_path):
    """Map a relative file path to the correct output subfolder."""
    rel_str = str(rel_path)
    # Try longest match first (more specific paths)
    for pattern, folder in sorted(FOLDER_MAP.items(), key=lambda x: -len(x[0])):
        if rel_str.startswith(pattern):
            return folder
    # Root-level files
    return ""


def strip_frontmatter(text):
    """Remove YAML frontmatter from markdown."""
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            return text[end + 3:].strip()
    return text


def setup_styles(doc):
    """Configure document styles for professional output."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Heading 1
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(20)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(16)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2E, 0x4A, 0x6E)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    if "Heading 3" in doc.styles:
        h3 = doc.styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0x4A, 0x6A, 0x8A)
        h3.paragraph_format.space_before = Pt(8)
        h3.paragraph_format.space_after = Pt(4)


def md_to_docx(md_text, output_path, source_path=""):
    """Convert markdown text to a formatted .docx file."""
    doc = Document()
    setup_styles(doc)

    # Add source footer
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    lines = md_text.split("\n")
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block — add as formatted paragraph
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    # Add light gray shading
                    from docx.oxml.ns import qn
                    shading = run._element.get_or_add_rPr()
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            if cells and not all(set(c) <= set("-: ") for c in cells):
                table_rows.append(cells)
            if not in_table:
                in_table = True
            i += 1
            # Check if next line is NOT a table row
            if i >= len(lines) or "|" not in lines[i] or not lines[i].strip().startswith("|"):
                # Render table
                if table_rows:
                    try:
                        num_cols = max(len(r) for r in table_rows)
                        table = doc.add_table(rows=len(table_rows), cols=num_cols)
                        table.style = "Light Grid Accent 1"
                        for ri, row_data in enumerate(table_rows):
                            for ci, cell_text in enumerate(row_data):
                                if ci < num_cols:
                                    cell = table.rows[ri].cells[ci]
                                    cell.text = cell_text
                                    for p in cell.paragraphs:
                                        p.style.font.size = Pt(9)
                    except Exception:
                        # Fallback: render as text
                        for row_data in table_rows:
                            doc.add_paragraph(" | ".join(row_data))
                table_rows = []
                in_table = False
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Horizontal rules
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Strip bold markers
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold text within bullets
            parts = re.split(r'(\*\*.+?\*\*)', line.strip()[2:])
            p.clear()
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+[\.\)]\s', line.strip()):
            text = re.sub(r'^\d+[\.\)]\s', '', line.strip())
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r'(\*\*.+?\*\*)', text)
            p.clear()
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.+?\*\*)', line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1

    # Add footer with source info
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run(f"Auto-generated from Claude Mainframe — {source_path}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def should_skip(rel_path):
    """Check if a file should be skipped."""
    rel_str = str(rel_path)
    for pattern in SKIP_PATTERNS:
        if pattern in rel_str:
            if pattern == "logs/" and INCLUDE_LOGS:
                continue
            return True
    return False


def main():
    # Clean output directory (skip .DS_Store and other system files)
    if os.path.exists(OUTPUT_DIR):
        for root, dirs, files in os.walk(OUTPUT_DIR, topdown=False):
            for f in files:
                if f == '.DS_Store':
                    continue
                try:
                    os.remove(os.path.join(root, f))
                except PermissionError:
                    pass
            for d in dirs:
                try:
                    os.rmdir(os.path.join(root, d))
                except (OSError, PermissionError):
                    pass
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    converted = 0
    skipped = 0
    errors = 0

    mainframe_path = Path(MAINFRAME_DIR)

    for md_file in sorted(mainframe_path.rglob("*.md")):
        rel_path = md_file.relative_to(mainframe_path)

        if should_skip(rel_path):
            skipped += 1
            continue

        try:
            # Read markdown
            with open(md_file, "r", encoding="utf-8") as f:
                content = f.read()

            # Strip frontmatter
            content = strip_frontmatter(content)

            if not content.strip():
                skipped += 1
                continue

            # Determine output folder
            output_folder = get_output_folder(rel_path)

            # Create output filename
            docx_name = md_file.stem + ".docx"

            if output_folder:
                output_path = os.path.join(OUTPUT_DIR, output_folder, docx_name)
            else:
                output_path = os.path.join(OUTPUT_DIR, docx_name)

            # Convert
            md_to_docx(content, output_path, str(rel_path))
            converted += 1
            print(f"  ✅ {output_folder}/{docx_name}" if output_folder else f"  ✅ {docx_name}")

        except Exception as e:
            errors += 1
            print(f"  ❌ {rel_path}: {e}")

    print(f"\n{'='*50}")
    print(f"Converted: {converted} files")
    print(f"Skipped: {skipped} files (logs, internal)")
    print(f"Errors: {errors} files")
    print(f"Output: {OUTPUT_DIR}")
    print(f"{'='*50}")


if __name__ == "__main__":
    main()
