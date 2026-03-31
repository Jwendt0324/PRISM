#!/usr/bin/env python3
"""
[Methodology Partner] Article QA Audit Script v2.0

Scans .docx files for BLOCK and WARN tier violations per the [Methodology Partner]
Article Quality Gate (canon: 03-article-guidelines.md).

Usage:
    python3 qa_audit.py <folder_path> [--fix-contractions] [--verbose]

Requirements:
    pip install python-docx

Output:
    - Per-article violation report (BLOCK / WARN)
    - Summary with pass/fail counts
    - Optional: auto-fix contractions (with artifact scan)
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# =============================================================================
# CANONICAL BANNED LISTS — Keep in sync with 03-article-guidelines.md
# =============================================================================

BANNED_WORDS = [
    "delve", "landscape", "realm", "paradigm", "synergy", "game-changer",
    "revolutionize", "cutting-edge", "harness", "utilize", "facilitate",
    "streamline", "robust", "foster", "spearhead", "holistic", "ecosystem",
    "empower", "pivot", "disrupt", "beacon",
]

# "leverage" is only banned as a verb. We check context separately.
LEVERAGE_VERB_PATTERNS = [
    r"\bleverage\s+(your|their|the|our|its|this|that|these|those)\b",
    r"\bleveraging\b",
    r"\bleveraged\s+(by|the|their|our)\b",
    r"\bto\s+leverage\b",
]

BANNED_AI_PATTERNS = [
    r"[Ii]n today'?s\s+\w+",
    r"[Ii]t'?s important to note",
    r"[Ww]hether you'?re a\b.*?\bor a\b",
    r"[Aa]t the end of the day",
    r"[Ii]n the world of",
    r"\b\w+\s+is not just about",
    r"[Ww]hen it comes to",
    r"[Ii]n conclusion",
    r"[Aa]s we navigate",
    r"[Nn]ot because\b.*?\bbut because",
    r"[Tt]his is where\b.*?\bcomes? in",
    r"[Ll]et that sink in",
    r"[Ff]ull stop\.\s*$",
]

BANNED_SALESY = [
    r"limited time",
    r"what are you waiting for",
    r"act now",
    r"don'?t miss out",
    r"sign up today",
]

# =============================================================================
# WARN-TIER CHECKS
# =============================================================================

EM_DASH_PATTERNS = [
    "\u2014",       # em dash
    "\u2013",       # en dash
    " -- ",         # double hyphen as em dash
]

CONTRACTION_MAP = {
    "do not": "don't",
    "is not": "isn't",
    "did not": "didn't",
    "does not": "doesn't",
    "was not": "wasn't",
    "has not": "hasn't",
    "cannot": "can't",
    "will not": "won't",
    "would not": "wouldn't",
    "could not": "couldn't",
    "should not": "shouldn't",
    "that is": "that's",
    "it is": "it's",
    "there is": "there's",
    "I am": "I'm",
    "you are": "you're",
    "we are": "we're",
    "they are": "they're",
    "he is": "he's",
    "she is": "she's",
    "I have": "I've",
    "we have": "we've",
    "you have": "you've",
    "I would": "I'd",
    "you would": "you'd",
    "I will": "I'll",
    "you will": "you'll",
    "we will": "we'll",
    "let us": "let's",
}

# Contractions that create broken artifacts — exclude from auto-fix
CONTRACTION_EXCLUSIONS = [
    # "does not" before words starting with "th" → "doesn'thing"
    (r"does not\s*(?=\b(?:th|n)\w)", "does not"),
    # "you have" followed by a bare noun (not "got", "been", "to")
    (r"you have\s+(?!got\b|been\b|to\b|had\b|just\b|already\b|never\b|always\b|ever\b)", None),
]

# Post-contraction artifact patterns to scan for
CONTRACTION_ARTIFACTS = [
    r"doesn'th",
    r"doesn'tn",
    r"you've\s+a\s+",  # British-sounding "you've a pile"
    r"(?<!\w)i\s+(?=[a-z])",  # lowercase "i " at sentence start
    r"it'self",
    r"n't\w{3,}",  # contraction fused with next word
]


def extract_text(docx_path: str) -> str:
    """Extract full text from a .docx file."""
    doc = Document(docx_path)
    return "\n".join(para.text for para in doc.paragraphs)


def extract_paragraphs(docx_path: str) -> list[str]:
    """Extract paragraphs as a list."""
    doc = Document(docx_path)
    return [para.text for para in doc.paragraphs if para.text.strip()]


def count_words(text: str) -> int:
    return len(text.split())


def count_contractions(text: str) -> int:
    """Count common English contractions in text."""
    contraction_pattern = r"\b\w+'(?:t|s|re|ve|ll|d|m)\b"
    return len(re.findall(contraction_pattern, text, re.IGNORECASE))


def find_internal_links(text: str) -> list[str]:
    """Find markdown-style or HTML links."""
    md_links = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', text)
    html_links = re.findall(r'href=["\']([^"\']+)["\']', text)
    return md_links + [(url,) for url in html_links]


def check_rhetorical_questions(paragraphs: list[str]) -> list[dict]:
    """Find stacked rhetorical questions (2+ in a row)."""
    violations = []
    consecutive_q = 0
    for para in paragraphs:
        sentences = re.split(r'[.!]\s+', para)
        for sent in sentences:
            if sent.strip().endswith('?'):
                consecutive_q += 1
                if consecutive_q >= 2:
                    violations.append({
                        "tier": "WARN",
                        "check": "Stacked rhetorical questions",
                        "detail": sent.strip()[:80],
                    })
            else:
                consecutive_q = 0
    return violations


def audit_article(docx_path: str, verbose: bool = False) -> dict:
    """Run full QA audit on a single .docx article."""
    text = extract_text(docx_path)
    paragraphs = extract_paragraphs(docx_path)
    text_lower = text.lower()
    word_count = count_words(text)
    violations = []

    # --- BLOCK: Banned words ---
    for word in BANNED_WORDS:
        matches = list(re.finditer(rf'\b{re.escape(word)}\b', text, re.IGNORECASE))
        for m in matches:
            # Get surrounding context
            start = max(0, m.start() - 30)
            end = min(len(text), m.end() + 30)
            context = text[start:end].replace('\n', ' ')
            violations.append({
                "tier": "BLOCK",
                "check": f"Banned word: {word}",
                "detail": f"...{context}...",
            })

    # --- BLOCK: Leverage as verb ---
    for pattern in LEVERAGE_VERB_PATTERNS:
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        for m in matches:
            start = max(0, m.start() - 30)
            end = min(len(text), m.end() + 30)
            context = text[start:end].replace('\n', ' ')
            violations.append({
                "tier": "BLOCK",
                "check": "Banned word: leverage (as verb)",
                "detail": f"...{context}...",
            })

    # --- BLOCK: Banned AI patterns ---
    for pattern in BANNED_AI_PATTERNS:
        matches = list(re.finditer(pattern, text))
        for m in matches:
            start = max(0, m.start() - 20)
            end = min(len(text), m.end() + 20)
            context = text[start:end].replace('\n', ' ')
            violations.append({
                "tier": "BLOCK",
                "check": "Banned AI pattern",
                "detail": f"...{context}...",
            })

    # --- BLOCK: Banned salesy language ---
    for pattern in BANNED_SALESY:
        # Check sentence context for "limited time" to avoid false positives
        if pattern == "limited time":
            for m in re.finditer(r'limited time', text, re.IGNORECASE):
                start = max(0, m.start() - 50)
                end = min(len(text), m.end() + 50)
                context = text[start:end].lower()
                # Skip if it's "unlimited time" or other false positives
                if "unlimited" in context:
                    continue
                violations.append({
                    "tier": "BLOCK",
                    "check": "Banned salesy language",
                    "detail": f"...{context.replace(chr(10), ' ')}...",
                })
        else:
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            for m in matches:
                start = max(0, m.start() - 30)
                end = min(len(text), m.end() + 30)
                context = text[start:end].replace('\n', ' ')
                violations.append({
                    "tier": "BLOCK",
                    "check": "Banned salesy language",
                    "detail": f"...{context}...",
                })

    # --- WARN: Em dashes ---
    for dash in EM_DASH_PATTERNS:
        count = text.count(dash)
        if count > 0:
            violations.append({
                "tier": "WARN",
                "check": f"Em/en dashes found: {count}",
                "detail": "Zero allowed unless voice profile permits",
            })

    # --- WARN: Word count ---
    if word_count < 200:
        violations.append({
            "tier": "WARN",
            "check": f"Word count: {word_count}",
            "detail": "Below 200-word minimum",
        })

    # --- WARN: Zero contractions ---
    contraction_count = count_contractions(text)
    if contraction_count == 0 and word_count >= 500:
        violations.append({
            "tier": "WARN",
            "check": "Zero contractions in 500+ word article",
            "detail": f"{word_count} words, 0 contractions — sounds robotic",
        })

    # --- WARN: Internal links ---
    links = find_internal_links(text)
    if len(links) < 2:
        violations.append({
            "tier": "WARN",
            "check": f"Internal links: {len(links)}",
            "detail": "Minimum 2 internal links required",
        })

    # --- WARN: Text walls (paragraphs over 8 lines) ---
    for i, para in enumerate(paragraphs):
        line_count = len(para.split('\n'))
        # Approximate: also check word count as proxy for visual length
        if count_words(para) > 150:
            violations.append({
                "tier": "WARN",
                "check": f"Text wall: paragraph {i+1}",
                "detail": f"{count_words(para)} words in single paragraph",
            })

    # --- WARN: Stacked rhetorical questions ---
    violations.extend(check_rhetorical_questions(paragraphs))

    # Compile result
    block_count = sum(1 for v in violations if v["tier"] == "BLOCK")
    warn_count = sum(1 for v in violations if v["tier"] == "WARN")

    return {
        "file": os.path.basename(docx_path),
        "path": docx_path,
        "word_count": word_count,
        "contraction_count": contraction_count,
        "link_count": len(links),
        "violations": violations,
        "block_count": block_count,
        "warn_count": warn_count,
        "status": "FAIL" if block_count > 0 else ("WARN" if warn_count > 0 else "PASS"),
    }


def fix_contractions(docx_path: str, dry_run: bool = False) -> dict:
    """Apply contraction fixes to a .docx file with artifact protection."""
    doc = Document(docx_path)
    changes = []
    artifact_warnings = []

    for para in doc.paragraphs:
        original = para.text

        # Check exclusions first
        for exc_pattern, _ in CONTRACTION_EXCLUSIONS:
            if re.search(exc_pattern, para.text, re.IGNORECASE):
                # Skip this specific contraction for this paragraph
                pass

        # Apply contractions
        new_text = para.text
        for formal, contracted in CONTRACTION_MAP.items():
            # Check exclusions before replacing
            skip = False
            for exc_pattern, _ in CONTRACTION_EXCLUSIONS:
                if formal.lower() in exc_pattern.lower() and re.search(exc_pattern, new_text, re.IGNORECASE):
                    skip = True
                    break
            if not skip:
                new_text = re.sub(
                    rf'\b{re.escape(formal)}\b',
                    contracted,
                    new_text,
                    flags=re.IGNORECASE,
                )

        # Artifact scan
        for artifact_pattern in CONTRACTION_ARTIFACTS:
            if re.search(artifact_pattern, new_text):
                artifact_warnings.append({
                    "file": os.path.basename(docx_path),
                    "original": original[:80],
                    "broken": new_text[:80],
                    "pattern": artifact_pattern,
                })
                new_text = original  # Revert this paragraph
                break

        if new_text != original:
            changes.append({
                "original": original[:80],
                "fixed": new_text[:80],
            })
            if not dry_run:
                para.text = new_text

    if not dry_run and changes:
        doc.save(docx_path)

    return {
        "file": os.path.basename(docx_path),
        "changes": len(changes),
        "artifacts_caught": len(artifact_warnings),
        "artifact_details": artifact_warnings,
    }


def main():
    parser = argparse.ArgumentParser(
        description="[Methodology Partner] Article QA Audit — scan .docx files for quality violations"
    )
    parser.add_argument("folder", help="Path to folder containing .docx files")
    parser.add_argument("--fix-contractions", action="store_true",
                        help="Auto-fix contractions (with artifact protection)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Show what --fix-contractions would change without modifying files")
    parser.add_argument("--verbose", "-v", action="store_true",
                        help="Show all violation details")
    args = parser.parse_args()

    folder = Path(args.folder)
    if not folder.is_dir():
        print(f"ERROR: {folder} is not a directory")
        sys.exit(1)

    docx_files = sorted(folder.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {folder}")
        sys.exit(0)

    print(f"\n{'='*70}")
    print(f"  [Methodology Partner] QA Audit — {len(docx_files)} articles")
    print(f"{'='*70}\n")

    # --- Run contraction fixes first if requested ---
    if args.fix_contractions:
        print("PHASE 1: Fixing contractions...\n")
        total_changes = 0
        total_artifacts = 0
        for docx_file in docx_files:
            result = fix_contractions(str(docx_file), dry_run=args.dry_run)
            if result["changes"] > 0 or result["artifacts_caught"] > 0:
                prefix = "[DRY RUN] " if args.dry_run else ""
                print(f"  {prefix}{result['file']}: {result['changes']} contractions fixed, "
                      f"{result['artifacts_caught']} artifacts caught")
                total_changes += result["changes"]
                total_artifacts += result["artifacts_caught"]
                for art in result["artifact_details"]:
                    print(f"    ⚠ ARTIFACT: {art['broken']}")

        print(f"\n  Total: {total_changes} contractions fixed, {total_artifacts} artifacts caught\n")
        print(f"{'='*70}\n")

    # --- Run audit ---
    print("QA AUDIT RESULTS\n")

    results = []
    for docx_file in docx_files:
        result = audit_article(str(docx_file), verbose=args.verbose)
        results.append(result)

        # Status indicator
        if result["status"] == "PASS":
            indicator = "✓ PASS"
        elif result["status"] == "WARN":
            indicator = "⚠ WARN"
        else:
            indicator = "✗ FAIL"

        print(f"  {indicator}  {result['file']}")
        print(f"         Words: {result['word_count']} | "
              f"Contractions: {result['contraction_count']} | "
              f"Links: {result['link_count']} | "
              f"BLOCK: {result['block_count']} | "
              f"WARN: {result['warn_count']}")

        if args.verbose or result["block_count"] > 0:
            for v in result["violations"]:
                tier_mark = "✗" if v["tier"] == "BLOCK" else "⚠"
                print(f"         {tier_mark} [{v['tier']}] {v['check']}")
                print(f"           {v['detail']}")
        print()

    # --- Summary ---
    pass_count = sum(1 for r in results if r["status"] == "PASS")
    warn_count = sum(1 for r in results if r["status"] == "WARN")
    fail_count = sum(1 for r in results if r["status"] == "FAIL")
    total_blocks = sum(r["block_count"] for r in results)
    total_warns = sum(r["warn_count"] for r in results)

    print(f"{'='*70}")
    print(f"  SUMMARY: {len(results)} articles scanned")
    print(f"  ✓ PASS: {pass_count} | ⚠ WARN: {warn_count} | ✗ FAIL: {fail_count}")
    print(f"  Total BLOCK violations: {total_blocks}")
    print(f"  Total WARN violations: {total_warns}")
    print(f"{'='*70}")

    if fail_count > 0:
        print(f"\n  {fail_count} article(s) have BLOCK violations and CANNOT be published.")
        print("  Run with --verbose to see all violation details.")
        print("  Fix violations, then re-run until 0 BLOCK.\n")
        sys.exit(1)
    elif warn_count > 0:
        print(f"\n  {warn_count} article(s) have WARN violations. Fix or document.")
        sys.exit(0)
    else:
        print(f"\n  All articles passed. Ready for human review.\n")
        sys.exit(0)


if __name__ == "__main__":
    main()
