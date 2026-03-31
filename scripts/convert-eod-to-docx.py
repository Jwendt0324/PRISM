#!/usr/bin/env python3
"""
PRISM EOD → DOCX Converter (Basecamp-optimized)
Converts an EOD markdown report to a .docx formatted for clean copy-paste into Basecamp.

Basecamp's Trix editor supports: bold, italic, links, bullets, numbered lists,
headings (H1 only), blockquotes, and attachments. No tables, no code blocks,
no custom fonts/colors, no nested lists.

Usage: python3 convert-eod-to-docx.py [YYYY-MM-DD]
       Defaults to today's date if no argument given.
"""

import os
import sys
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Inches


PRISM_DIR = os.path.expanduser("~/Documents/Claude/PRISM")
REPORTS_DIR = os.path.join(PRISM_DIR, "logs/reports")
EODS_DIR = os.path.join(REPORTS_DIR, "eods")


def setup_styles(doc):
    """Minimal styling — Basecamp strips most formatting on paste."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for level_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        if level_name in doc.styles:
            h = doc.styles[level_name]
            h.font.name = "Calibri"
            h.font.size = Pt(size)
            h.font.bold = True


def parse_inline(paragraph, text):
    """Parse bold and inline code from markdown into DOCX runs."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            # No monospace in Basecamp — just bold it
            run = paragraph.add_run(part[1:-1])
            run.bold = True
        elif part.startswith("["):
            m = re.match(r'\[(.+?)\]\((.+?)\)', part)
            if m:
                paragraph.add_run(m.group(1))
            else:
                paragraph.add_run(part)
        elif part:
            paragraph.add_run(part)


def table_to_list(rows):
    """Convert markdown table rows into 'Label: Value' pairs for Basecamp."""
    if not rows:
        return []
    # First row is header
    header = rows[0] if rows else []
    result = []
    for row in rows[1:]:
        if len(row) >= 2 and len(header) >= 2:
            result.append(f"{row[0]}: {row[1]}")
        elif len(row) >= 1:
            result.append(" — ".join(row))
    return result


def md_to_docx(md_text, output_path):
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    lines = md_text.split("\n")
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code blocks — render as plain bold text (no monospace in Basecamp)
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.bold = True
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables — convert to bold label: value bullet list
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            # Skip separator rows (---|---)
            if cells and not all(set(c) <= set("-: ") for c in cells):
                table_rows.append(cells)
            if not in_table:
                in_table = True
            i += 1
            # Check if table ended
            if i >= len(lines) or "|" not in lines[i] or not lines[i].strip().startswith("|"):
                if table_rows:
                    items = table_to_list(table_rows)
                    for item in items:
                        p = doc.add_paragraph(style="List Bullet")
                        # Bold the label part (before the colon)
                        if ": " in item:
                            label, value = item.split(": ", 1)
                            run = p.add_run(label + ": ")
                            run.bold = True
                            parse_inline(p, value)
                        else:
                            parse_inline(p, item)
                table_rows = []
                in_table = False
            continue

        # Empty lines — add spacing
        if not line.strip():
            i += 1
            continue

        # Headings — Basecamp only really supports H1, but bold H2/H3 work on paste
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1:
                doc.add_heading(text, level=1)
            else:
                # H2/H3 as bold paragraphs with size distinction
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14 if level == 2 else 12)
            i += 1
            continue

        # Horizontal rules — just add a blank line (Basecamp has no HR)
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # Checkboxes — plain bullets with text markers
        if line.strip().startswith("- [ ] ") or line.strip().startswith("- [x] "):
            checked = line.strip().startswith("- [x]")
            text = line.strip()[6:]
            p = doc.add_paragraph(style="List Bullet")
            if checked:
                run = p.add_run("[DONE] ")
                run.bold = True
            else:
                run = p.add_run("[TODO] ")
                run.bold = True
            parse_inline(p, text)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\s*(\d+)\.\s+(.+)$', line)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, text)
            i += 1
            continue

        # Italic footer
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.italic = True
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, line.strip())
        i += 1

    doc.save(str(output_path))
    return output_path


def main():
    target_date = sys.argv[1] if len(sys.argv) > 1 else date.today().strftime("%Y-%m-%d")

    md_path = os.path.join(REPORTS_DIR, f"eod-{target_date}.md")
    if not os.path.exists(md_path):
        print(f"ERROR: No EOD report found at {md_path}")
        sys.exit(1)

    os.makedirs(EODS_DIR, exist_ok=True)

    with open(md_path, "r") as f:
        md_text = f.read()

    # Strip YAML frontmatter if present
    if md_text.startswith("---"):
        end = md_text.find("---", 3)
        if end != -1:
            md_text = md_text[end + 3:].strip()

    output_path = os.path.join(EODS_DIR, f"eod-{target_date}.docx")
    md_to_docx(md_text, output_path)
    print(f"DOCX written to {output_path}")


if __name__ == "__main__":
    main()
