#!/usr/bin/env python3
"""
Converts all .md SOPs in the Mainframe to readable .docx files.
The .md files stay for Claude. The .docx files are for Jack.
Output: sops/_readable/ folder with .docx versions.
"""
import os, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOPS_DIR = Path(__file__).parent
OUTPUT_DIR = SOPS_DIR / "_readable"
OUTPUT_DIR.mkdir(exist_ok=True)

def md_to_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    lines = md_path.read_text(encoding='utf-8').splitlines()

    # Skip frontmatter
    i = 0
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1  # skip closing ---

    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        i += 1

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - write accumulated code
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Empty lines
        if not line.strip():
            continue

        # H1
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
            continue

        # H2
        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            continue

        # H3
        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Checklist items
        if line.strip().startswith('- [ ]'):
            text = line.strip()[5:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, "☐ " + text)
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue

        # Numbered items
        m = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            continue

        # Sub-items (a, b, c)
        m2 = re.match(r'^([a-z])\.\s+(.*)', line.strip())
        if m2:
            text = m2.group(1) + ". " + m2.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.75)
            add_formatted_text(p, text)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())

    doc.save(str(docx_path))


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, code, links) into runs."""
    # Split on bold (**text**) and code (`text`) markers
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
        elif part.startswith('['):
            # Link: [text](url) - just show the text
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x00, 0x66, 0xcc)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def main():
    converted = 0
    for root, dirs, files in os.walk(SOPS_DIR):
        # Skip output dir and templates
        if '_readable' in root:
            continue
        for f in sorted(files):
            if f.endswith('.md') and f != 'INDEX.md':
                md_path = Path(root) / f
                docx_name = f.replace('.md', '.docx')
                docx_path = OUTPUT_DIR / docx_name
                try:
                    md_to_docx(md_path, docx_path)
                    print(f"  ✅ {docx_name}")
                    converted += 1
                except Exception as e:
                    print(f"  ❌ {f}: {e}")

    print(f"\nConverted {converted} SOPs to docx")
    print(f"Output: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
